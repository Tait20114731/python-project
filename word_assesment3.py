# -*- coding: utf-8 -*-
"""
Created on Mon Nov  4 19:44:32 2024

@author: Tatenda Marimo 20114731
@description: program creates a new Document, adds headings, pragpraphs and a
picture then finaly reads and saves the Final document.
@version: 1.0
"""
from docx import Document
from docx.shared import Inches

# Creates a new Word document and returns the Document object.
def create_document():
    document = Document()
    return document

# Adds headings to the document.
def add_headings(document):
    # Adding the main heading (Level 0)
    document.add_heading('Building Documents Using Python', level=0)
    # Adding Header 1 (Level 1)
    document.add_heading('This is Header 1 as a Main Heading', level=1)
    # Adding Header 3 (Level 3)
    document.add_heading('Activities', level=3)

# Adds paragraphs and lists under the headings.
def add_paragraphs(document):
    # Adding a simple paragraph under Header 1
    document.add_paragraph('This is a simple and first paragraph.')
    # Adding a numbered list item under Header 3
    document.add_paragraph('Swimming', style='List Number')
    # Adding a numbered list item under Header 3
    document.add_paragraph('Running', style='List Number')
    # Adding another heading for the list
    document.add_heading('Things to Do', level=3)
    # Adding bullet list items under "Things to Do"
    document.add_paragraph('Excessive eating', style='List Bullet')
    document.add_paragraph('Excessive swimming', style='List Bullet')

# Adds a picture to the document.
def add_picture(document, image_path):

    try:
        # Adding the picture with specified dimensions
        document.add_picture(image_path, width=Inches(12.2), height=Inches(8.7))
    except Exception as e:
        print(f"Error adding picture: {e}")

# Retrieves and returns the full text from the specified Word document.
def get_full_text(doc_filename):
    document = Document(doc_filename)
    final_text = []
    for paragraph in document.paragraphs:
        final_text.append(paragraph.text)
    return '\n'.join(final_text)

# Saves the document to the specified filename.
def save_document(document, filename):
    document.save(filename)

def main():
    # Create a new document
    doc = create_document()
    
    # Add headings
    add_headings(doc)
    
    # Add paragraphs and lists
    add_paragraphs(doc)
    
    # Add a picture
    image_path = 'screenshot.png'  
    add_picture(doc, image_path)
    
    # Save the document
    doc_filename = 'Final_doc.docx'
    save_document(doc, doc_filename)
    
    # Retrieve and print the full text from the document
    full_text = get_full_text('fullText.docx')
    print("Full text from the document:")
    print(full_text)

if __name__ == '__main__':
    main()
