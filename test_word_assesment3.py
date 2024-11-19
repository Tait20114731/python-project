# -*- coding: utf-8 -*-
"""
Created on Tue Nov 19 16:14:13 2024

@author: Tatenda Marimo
@description: Unit testing for word_assesment3 program. This module includes tests for each function,
ensuring word document can be created, add headings, add paragraphs, add picture, read text and save document.
@version: 1.0
"""

import unittest
import os
from PIL import Image
from docx.document import Document as DocumentClass
from word_assesment3 import create_document, add_headings, add_paragraphs, add_picture, get_full_text, save_document


class TestWordAssesment3(unittest.TestCase):

    def setUp(self):
        # Setup sample document file paths
        self.doc_filename = 'test_doc.docx'
        self.image_path = 'sample_image.png'

        # Create a valid sample image file for testing
        image = Image.new('RGB', (100, 100), color='red')
        image.save(self.image_path)

    def tearDown(self):
        # Clean up created files after tests
        files_to_remove = [self.doc_filename, self.image_path]
        for file in files_to_remove:
            if os.path.exists(file):
                os.remove(file)

    def test_create_document(self):
        # Test creating a new document
        doc = create_document()
        self.assertIsInstance(doc, DocumentClass, "The created object is not an instance of Document.")

    def test_add_headings(self):
        # Test adding headings to the document
        doc = create_document()
        add_headings(doc)
        self.assertEqual(len(doc.paragraphs), 3, "Headings were not added correctly.")
        self.assertEqual(doc.paragraphs[0].text, 'Building Documents Using Python')
        self.assertEqual(doc.paragraphs[1].text, 'This is Header 1 as a Main Heading')
        self.assertEqual(doc.paragraphs[2].text, 'Activities')

    def test_add_paragraphs(self):
        # Test adding paragraphs and lists to the document
        doc = create_document()
        add_headings(doc)
        add_paragraphs(doc)
        expected_paragraph_count = 9  # 3 headings + 6 paragraphs/lists
        self.assertEqual(len(doc.paragraphs), expected_paragraph_count, "Paragraphs and lists were not added correctly.")

    def test_add_picture(self):
        # Test adding a picture to the document
        doc = create_document()
        try:
            add_picture(doc, self.image_path)
        except Exception as e:
            self.fail(f"add_picture raised an unexpected exception: {e}")
        self.assertEqual(len(doc.inline_shapes), 1, "Picture was not added correctly.")

    def test_save_document(self):
        # Test saving the document
        doc = create_document()
        save_document(doc, self.doc_filename)
        self.assertTrue(os.path.exists(self.doc_filename), "The document was not saved correctly.")

    def test_get_full_text(self):
        # Test retrieving full text from the document
        doc = create_document()
        add_headings(doc)
        add_paragraphs(doc)
        save_document(doc, self.doc_filename)
        full_text = get_full_text(self.doc_filename)
        expected_text = (
            'Building Documents Using Python\n'
            'This is Header 1 as a Main Heading\n'
            'Activities\n'
            'This is a simple and first paragraph.\n'
            'Swimming\n'
            'Running\n'
            'Things to Do\n'
            'Excessive eating\n'
            'Excessive swimming'
        )
        # Strip trailing whitespace for comparison
        self.assertEqual(full_text.strip(), expected_text.strip(), "The retrieved full text is not correct.")

if __name__ == '__main__':
    print("Present", __name__)
    unittest.main()
